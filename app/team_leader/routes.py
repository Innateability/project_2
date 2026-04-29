from flask import Blueprint,url_for,render_template,redirect,flash,get_flashed_messages,request,session,abort
from app.models import Authentication,Objective,TeamLeader,Employee,Review,ObjectiveBatch,AdminObjective,Feedback,TeamLeaderFeedback,TeamLeaderFeedback,EmployeeEmail,Messages,ReviewOpenObjective,AuthReviewed,Department
from app import db
from sqlalchemy.exc import IntegrityError
from app.utils import login_required,team_leader_required
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


team_leader_bp = Blueprint("team_leader",__name__,url_prefix="/Team-Leader")

@team_leader_bp.route("/",methods=["GET"])
@login_required
@team_leader_required
def home():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("home.html",state="team_leader",role="team_leader",auth_name=auth_name,auth=auth)

@team_leader_bp.route("/logout",methods=["POST","GET"])
@login_required
@team_leader_required
def logout():
    auth_name = Authentication.query.get(session["user_id"]).name
    if request.method == "POST":
        session.pop("role",None)
        session.pop("user_id",None)
        user_id = session.get("user_id")
        auth = Authentication.query.get(user_id)
        return redirect(url_for("base.login",auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("logout.html",role="team_leader",state="team_leader",auth=auth)

@team_leader_bp.route("/delete_account",methods=["POST","GET"])
@login_required
@team_leader_required
def delete_account():
    if request.method == "POST":
        user_id = session.get("user_id")
        user = Authentication.query.filter_by(id=user_id).first()
        session.pop("role",None)
        session.pop("user_id",None)
        db.session.delete(user)
        try:
            db.session.commit()
        except IntegrityError:
            flash("username does not exist","error")
            db.session.rollback()
            return redirect(url_for("base.signup"))
        flash("account deleted succesfully","success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("base.login",auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("delete_account.html",state="team_leader",role="team_leader",auth=auth)

@team_leader_bp.route("/add_recipients/<int:batch_id>")
@login_required
@team_leader_required
def recipients(batch_id):
    current_auth = Authentication.query.get(session.get("user_id"))
    team_leader = current_auth.team_leader
    dept_id = team_leader.department_id
    
    existing_ids = db.session.query(Objective.assigned_to_id)\
        .filter(Objective.batch_id == batch_id)
    
    names = db.session.query(Authentication)\
        .join(Employee, Authentication.id == Employee.authentication_id)\
            .filter(
                Employee.department_id == dept_id,
                Authentication.id != current_auth.id,
                ~Authentication.id.in_(existing_ids))\
                    .order_by(Authentication.name.asc())\
                        .all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("team_leader_recipients.html",names_auth=names,state="team_leader",auth=auth,role="team_leader",batch_id=batch_id)

@team_leader_bp.route("/select_member")
@login_required
@team_leader_required
def select_member():
    current_user_id = session["user_id"]
    auth_id = Authentication.query.get(current_user_id).id
    current_user_department_id = (Authentication.query.get(current_user_id).team_leader.department_id)
    names_auth = (Authentication.query.join(Authentication.employee).filter(Authentication.id != current_user_id,Employee.department_id == current_user_department_id).order_by(Authentication.name.asc()).all())
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("team_leader_select_member.html",names_auth=names_auth,admin_id=current_user_id,state="team_leader",auth=auth,role="team_leader")

@team_leader_bp.route("/select_batch",methods=["POST","GET"])
@login_required
@team_leader_required
def select_batch():
    batches = ObjectiveBatch.query.all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("select_batch.html",auth=auth,role="team_leader",batches=batches,state="team_leader")

@team_leader_bp.route("/choose_batch",methods=["POST","GET"])
@login_required
@team_leader_required
def select_batch_for_open_batch():
    batches = ObjectiveBatch.query.all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("select_batch_for_open_batch.html",auth=auth,role="team_leader",batches=batches,state="team_leader")

@team_leader_bp.route("/reports/<auth_id>/<batch_id>",methods=["POST","GET"])
@login_required
@team_leader_required
def reports(auth_id,batch_id):
    auth = Authentication.query.get(session.get("user_id"))
    grouped = defaultdict(list)
    reports = defaultdict(list)
    batch = ObjectiveBatch.query.get(batch_id)
    active = batch.active
    team_leader = TeamLeader.query.filter_by(authentication_id=auth.id).first()
    team_leader_id = team_leader.id
    objectives = Objective.query.filter_by(batch_id=batch_id,assigned_by_id=team_leader_id).all()
    your_objectives = AdminObjective.query.filter_by(batch_id=batch_id,assigned_to_id=auth.id).all()
    auths = Authentication.query\
    .join(Employee, Employee.authentication_id == Authentication.id)\
    .join(Department, Department.id == Employee.department_id)\
    .join(TeamLeader, TeamLeader.department_id == Department.id)\
    .filter(
        Authentication.id != auth_id,
        TeamLeader.id == team_leader_id
    ).all()
    grouped_by_score = defaultdict(int)
    grouped_by_your_score = defaultdict(int)
    for auth in auths:
        open_total_weighted = 0
        objs = AdminObjective.query.filter_by(batch_id=batch_id, assigned_to_id=auth.id).all() + Objective.query.filter_by(batch_id=batch_id, assigned_to_id=auth.id).all()
        if objs:
            for obj in objs:
                if obj.open_objectives_review:
                    open_total_weighted += obj.open_objectives_review.weighted_score
        grouped_by_score[auth.id] += open_total_weighted
    for obj in objectives:
        total_weighted = 0
        if hasattr(obj, "review") and obj.review and obj.review.weighted_score:
            total_weighted += obj.review.weighted_score
        if total_weighted:
            key = (obj.batch.title, obj.batch.year, obj.assigned_to)
            op_total_weighted = total_weighted + grouped_by_score[obj.assigned_to_id]
            grouped[key].append(op_total_weighted)            
    for key, op_total_weighted in grouped.items():
        title, year, assigned_to = key
        objs = Objective.query.filter_by(batch_id=batch_id,assigned_to_id=assigned_to.id).first()
        obj_id = objs.id
        mode = "t"
        name_obj = (assigned_to,obj_id,mode)
        reports[name_obj].append(sum(op_total_weighted))   
    your_report = defaultdict(list)
    your_grouped = defaultdict(list)
    
    open_total_weighted = 0
    auth = Authentication.query.get(session.get("user_id"))
    objs = AdminObjective.query.filter_by(batch_id=batch_id, assigned_to_id=auth.id).all()
    if objs:
        for obj in objs:
            if obj.open_objectives_review:
                open_total_weighted += obj.open_objectives_review.weighted_score             
        grouped_by_your_score[auth.id] += open_total_weighted
    for obj in your_objectives:
        total_weighted = 0
        if hasattr(obj, "admin_review") and obj.admin_review:
            total_weighted += obj.admin_review.weighted_score
        if total_weighted:
            key = (obj.batch.title, obj.batch.year, obj.assigned_to)
            ops_total_weighted = total_weighted + grouped_by_your_score[obj.assigned_to_id]
            your_grouped[key].append(ops_total_weighted)
    for key, ops_total_weighted in your_grouped.items():
        title, year, assigned_to = key
        objs = AdminObjective.query.filter_by(batch_id=batch_id,assigned_to_id=assigned_to.id).first()
        if objs:
            obj_id = obj.id
            mode = "a"
            name_obj = (assigned_to,obj_id,mode)
            your_report[name_obj].append(sum(ops_total_weighted))
    now = datetime.now()
    return render_template("reports.html",auth=auth,role="team_leader",state="team_leader",batch=batch,reports=reports,your_reports=your_report,active=active,now=now)

@team_leader_bp.route("/download-report-word/<int:objective_id>")
def download_report_word(objective_id):
    doc = Document()
    doc.add_heading("Objective Report", 0)
    doc.add_paragraph(f"Objective ID: {objective_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Message"
    hdr[1].text = "Status"
    hdr[2].text = "Timestamp"

    # add data rows
    messages = Messages.query.filter_by(objective_id=objective_id).all()
    for msg in messages:
        row = table.add_row().cells
        row[0].text = msg.message
        row[1].text = msg.status
        row[2].text = str(msg.timestamp)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@team_leader_bp.route("/download-report/<int:objective_id>")
def download_report(objective_id):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # write content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Objective Report")

    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, f"Objective ID: {objective_id}")
    p.drawString(50, height - 100, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p.save()
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")


@team_leader_bp.route("/objective_batches",methods=["POST","GET"])
@login_required
@team_leader_required
def objective_batches():
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    batches = ObjectiveBatch.query.order_by(ObjectiveBatch.year.desc()).all()
    now = datetime.now()
    return render_template("objective_batches.html",auth=auth,role="team_leader",state="team_leader",names_auth=names_auth,batches=batches,now=now)

@team_leader_bp.route("/add_objectives/<batch_id>",methods=["POST","GET"])
@login_required
@team_leader_required
def add_objective(batch_id):
    authen = Authentication.query.get(session["user_id"])
    team_leader = authen.team_leader if authen else None
    now = datetime.now()
    if request.method == "GET":
        
        recipient_ids = request.args.getlist("recipients[]")
        recipients = Authentication.query.filter(Authentication.id.in_(recipient_ids)).all()
        auth = Authentication.query.get(session.get("user_id"))
        batch = ObjectiveBatch.query.get(batch_id)
        active = batch.active
        print(active)
        return render_template("team_leader_add_objective.html", recipients=recipients,state="team_leader",auth=auth,role="team_leader",batch_id=batch_id,batch=batch,active=active,now=now)
    if request.method == "POST":
        recipient_ids = request.form.getlist("recipients[]")
        recipients = Authentication.query.filter(Authentication.id.in_(recipient_ids)).all()
        objectives = request.form.getlist("objectives[]")
        categories = request.form.getlist("categories[]")
        weights = request.form.getlist("weights[]")
        publics = request.form.getlist("public[]")
        score_range = 5

        batch = ObjectiveBatch.query.get(batch_id)

        for obj, cat, weight, public in zip(
            objectives, categories, weights, publics):
            if public == "True":
                private = False
            else:
                private = True
            for recipient in recipients:
                objective = Objective(
                    objective=obj,
                    category=cat,
                    weight=int(weight),
                    score_range=int(score_range),
                    assigned_to=recipient,         
                    assigned_by=team_leader,
                    batch=batch,
                    private=private
                )
                auth_id = recipient.id
                db.session.add(objective)

        db.session.commit()
        flash("Objectives created successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives",auth_id=auth.id,batch_id=batch_id))

@team_leader_bp.route("/edit_objectives/<int:objective_id>", methods=["POST", "GET"])
@login_required
@team_leader_required
def edit_objective(objective_id):
    team_leader_email = session.get("email")
    department_head = Authentication.query.filter_by(email=team_leader_email).first().team_leader
    obj = Objective.query.filter_by(id=objective_id).first()
    if obj.assigned_by != department_head:
        abort(403)
    now = datetime.now()
    if request.method == "POST":
        objective_text = request.form.get("objective")
        category = request.form.get("category")
        weight = int(request.form.get("weight"))
        obj = Objective.query.get(objective_id).first()
        batch = obj.batch
        auth_id = obj.assigned_to_id
        public = request.form.get("public")
        if public:
            private = False
        else:
            private = True
        obj.objective = objective_text
        obj.category = category
        obj.weight = weight
        obj.private = private
        db.session.commit()
        flash(f"Objective {batch.title} edited successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives",auth_id=auth_id,auth=auth))

    department = department_head.department
    employees = department.employees
    objective = Objective.query.filter_by(id=objective_id).first()
    active = objective.batch.active
    employee_names = [emp.authentication.name for emp in employees]
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("team_leader_edit_objective.html",role="team_leader",state="team_leader",employee_names=employee_names,objective=objective,Title="EDIT OBJECTIVES",auth=auth,active=active,now=now)

@team_leader_bp.route("/received_objectives")
@login_required
@team_leader_required
def received_objectives():
    team_leader_auth =  Authentication.query.get(session.get("user_id"))
    received_objectives = (AdminObjective.query.filter_by(assigned_to=team_leader_auth).all())
    team_leader_grouped = defaultdict(list)
    for obj in received_objectives:
        key = obj.batch.title
        team_leader_grouped[key].append(obj)
    auth = Authentication.query.get(session.get("user_id"))
    mode = "a"
    active = received_objectives[0].batch.active
    now = datetime.now()
    return render_template("team_leader_received_objectives.html",team_leader_grouped_objectives=team_leader_grouped,state='team_leader',auth=auth,role="team_leader",mode=mode,active=active,now=now)


@team_leader_bp.route("/objectives/<int:auth_id>/<int:batch_id>")
@login_required
@team_leader_required
def objectives(auth_id,batch_id):
    auth = Authentication.query.get(auth_id)
    batch = ObjectiveBatch.query.get(batch_id)
    team_leader_auth =  Authentication.query.get(session.get("user_id")).team_leader
    active = batch.active
    if auth.team_leader:
        mode = "See All"
        objectives = Objective.query.filter(Objective.batch_id==batch_id).all()
        grouped = defaultdict(list)
        grouped_by_title = defaultdict(list)
        title = batch.title
        for obj in objectives:
            key = (obj.batch,obj.assigned_by,obj.assigned_to)
            grouped[key] = obj
        grouped_by_person = defaultdict(list)
        for key, obj in grouped.items():
            grouped_by_person[obj.assigned_to.name].append(obj)
        username = auth.name
        active = batch.active
        now = datetime.now()
        return render_template("team_leader_assigned_objectives.html",grouped_objectives=grouped_by_person,state='team_leader',name=username,mode=mode,auth=auth,role="team_leader",batch=batch,title=title,active=active,now=now)

@team_leader_bp.route("/review_open_objective/<int:objective_id>/<mode>", methods=["POST", "GET"])
@login_required
@team_leader_required
def review_open_objective(objective_id,mode):
    auth_id = session.get("user_id")
    if mode == "t":
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=objective_id,auth_id=auth_id).all()
    objective = Objective.query.get(objective_id) if mode == "t" else AdminObjective.query.get(objective_id)
    now = datetime.now()
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(request.url,auth=auth)
        weighted_score = (score * objective.weight) / objective.score_range
        obj = objective.open_objectives_review
        batch = objective.batch
        if obj and obj.weighted_score:
            weighted_scores = obj.weighted_score
        else:
            weighted_scores = 0
        
        if obj and obj.number_reviews:
            number = obj.number_reviews
        else:
            number = 0
        new_number = number + 1
        new_weighted_score = ((weighted_scores * number) + weighted_score)/new_number
        if obj:
            obj.weighted_score = new_weighted_score
            obj.number_reviews = new_number
            obj.score = score
            obj.review = review_text
        else:
            obj = ReviewOpenObjective(
                review=review_text,
                score=score,
                weighted_score=new_weighted_score,
                number_reviews=new_number)
            objective.open_objectives_review = obj
        if isinstance(objective, Objective):
            auth_review = AuthReviewed(
                score=score,
                auth_id=auth_id,
                objective_id = objective_id)
        else:
            auth_review = AuthReviewed(
                score=score,
                auth_id=auth_id,
                admin_objective_id = objective_id)
        db.session.add(auth_review)
        db.session.commit()
        flash("reviewed successfully","success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.open_objectives_overview", objective_id=objective.id, auth=auth, mode=mode, batch_id=batch.id))
    auth = Authentication.query.get(session.get("user_id"))
    active = objective.batch.active
    return render_template("team_leader_review.html", objective=objective, role="team_leader", state="team_leader",auth=auth, mode=mode, auth_reviewed=auth_reviewed,active=active,now=now)


@team_leader_bp.route("/review/<int:objective_id>", methods=["POST", "GET"])
@login_required
@team_leader_required
def review_objective(objective_id):
    objective = Objective.query.get(objective_id)
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(request.url,auth=auth)
        weighted_score = (score * objective.weight) / objective.score_range
        r = Review(
            review=review_text,
            score=score,
            weighted_score=weighted_score
        )
        objective.review = r
        db.session.commit()
        flash("reviewed successfully","success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives_overview", objective_id=objective.id, auth=auth, mode="t"))
    auth = Authentication.query.get(session.get("user_id"))
    active = objective.batch.active
    now = datetime.now()
    return render_template("team_leader_review.html", objective=objective, role="team_leader", state="team_leader",auth=auth,active=active,now=now)

@team_leader_bp.route("/edit_review/<int:objective_id>", methods=["POST", "GET"])
@login_required
@team_leader_required
def edit_review(objective_id):
    objective = Objective.query.get(objective_id)
    review = objective.review
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(request.url,auth=auth)
        weighted_score = (score * objective.weight) / objective.score_range

        review.review = review_text
        review.score = score
        review.weighted_score = weighted_score
        db.session.commit()
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives_overview", objective_id=objective.id, auth=auth, mode="t"))
    auth = Authentication.query.get(session.get("user_id"))
    active = objective.batch.active
    now = datetime.now()
    return render_template("team_leader_edit_review.html", objective=objective,review=review,state="team_leader",auth=auth,role="team_leader",active=active,now=now)

@team_leader_bp.route("/open_objectives/<int:batch_id>",methods=["POST","GET"])
@login_required
@team_leader_required
def open_objectives(batch_id):
    auth = Authentication.query.get(session.get("user_id"))
    objs = defaultdict(list)
    open_objectives = AdminObjective.query.filter(AdminObjective.batch_id == batch_id,AdminObjective.private != True,AdminObjective.assigned_to_id != auth.id).all() + Objective.query.filter(Objective.batch_id == batch_id,Objective.private != True,Objective.assigned_to_id != auth.id).all()
    now = datetime.now()
    for open_objective in open_objectives:
        key = (open_objective.assigned_to,open_objective.batch)
        objs[key] = open_objective

    open_objective = []
    for key,obj in objs.items():
        if isinstance(obj, Objective):
            mode = "t"
        else:
            mode = "a"
        open_objective.append((obj, mode))
    batch = ObjectiveBatch.query.get(batch_id)
    active = batch.active
    return render_template("open_objectives.html",open_objectives=open_objective,state="team_leader",role="team_leader",batch=batch,now=now,active=active)

@team_leader_bp.route("/open_objectives_overview/<int:batch_id>/<int:objective_id>/<mode>",methods=["POST","GET"])
@login_required
@team_leader_required
def open_objectives_overview(batch_id,objective_id,mode):
    auth = Authentication.query.get(session["user_id"])
    auth_name = auth.name
    auth_id = auth.id

    if mode == "a":
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
        objective = AdminObjective.query.get(objective_id)
        batch = objective.batch
        employee = objective.assigned_to
        objectives = AdminObjective.query.filter_by(batch_id=batch_id, assigned_to_id=employee.id, private=False).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
        objective = Objective.query.get(objective_id)
        batch = objective.batch
        employee = objective.assigned_to
        objectives = Objective.query.filter_by(batch_id=batch_id, assigned_to_id=employee.id, private=False).all()
    total_weighted = 0
    objs = {}
    for obj in objectives:
        if mode == "a":
            auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=obj.id,auth_id=auth_id).first()
        else:
            auth_reviewed = AuthReviewed.query.filter_by(objective_id=obj.id,auth_id=auth_id).first()
        if hasattr(obj, "open_objectives_review"):  
            if obj.open_objectives_review:
                key = (obj,auth_reviewed)
                objs[key] = obj.open_objectives_review.weighted_score
                total_weighted += obj.open_objectives_review.weighted_score
        
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    active = ObjectiveBatch.query.get(batch_id).active
    print(active)
    return render_template("open_objectives_overview.html",employee=employee,objectives=objs,objective=objectives,total_weighted=total_weighted,batch=batch,state='team_leader',auth=auth,role="team_leader",now=now,mode=mode,auth_reviewed=auth_reviewed,active=active)

@team_leader_bp.route("/open_objective_overview/<int:objective_id>/<int:assigned_by_id>/<mode>", methods=["POST","GET"])
@login_required
@team_leader_required
def open_objective_overview(objective_id, assigned_by_id, mode):
    auth_id = session.get("user_id")
    if mode == "t":
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=objective_id,auth_id=auth_id).all()
    if request.method == "POST":
        message = request.form.get("message")
        if message:            
            now = datetime.now()
            new_message = Messages(message=message,status="team_leader",timestamp=now,admin_objective_id=objective_id)
            db.session.add(new_message)
            db.session.commit()
    message = Messages.query.filter_by(admin_objective_id=objective_id).delete()
    db.session.commit()
    messages = Messages.query.filter_by(admin_objective_id=objective_id).order_by(Messages.timestamp.asc()).all()
    assigned_by_auth = Authentication.query.get(assigned_by_id)
    if assigned_by_auth.role == "team_leader":
        assigned_by = assigned_by_auth.team_leader
        objective = Objective.query.get(objective_id)
        batch = objective.batch
    elif assigned_by_auth.role == "admin":
        assigned_by = assigned_by_auth.administrator
        objective = AdminObjective.query.get(objective_id)
        batch = objective.batch
    else:
        abort(404)
    title = batch.title
    year = batch.year
    employee = objective.assigned_to.name
    auth = Authentication.query.get(session.get("user_id"))
    active = batch.active
    now = datetime.now()
    return render_template(
        "open_objective_overview.html",
        year=year,
        title=title,
        objective=objective,
        employee=employee,
        role="team_leader",
        state="team_leader",
        messages=messages,
        mode=mode,
        auth = Authentication.query.get(session.get("user_id")),
        auth_reviewed=auth_reviewed,
        active=active,
        now=now)


@team_leader_bp.route("/objectives_overview/<mode>/<int:objective_id>")
@login_required
@team_leader_required
def objectives_overview(mode,objective_id):
    total_weighted = 0
    final_total_weighted = 0
    if mode == "a":
        objective = AdminObjective.query.get(objective_id)
        employee = objective.assigned_to
        batch = objective.batch
        objectives = AdminObjective.query.filter_by(batch_id=batch.id, assigned_to_id=employee.id).all()
        for obj in objectives:
            if obj.admin_review:
                total_weighted += obj.admin_review.weighted_score
                final_total_weighted += obj.admin_review.weighted_score
                auth = Authentication.query.get(session.get("user_id"))
    else:
        objective = Objective.query.get(objective_id)   
        employee = objective.assigned_to
        batch = objective.batch
        objectives = Objective.query.filter_by(batch_id=batch.id, assigned_to_id=employee.id).all()
        for obj in objectives:
            if obj.review:
                total_weighted += obj.review.weighted_score
                final_total_weighted += obj.review.weighted_score
    for obj in objectives:
        if obj.open_objectives_review:
            final_total_weighted += obj.open_objectives_review.weighted_score
    auth = Authentication.query.get(session.get("user_id"))
    active = batch.active
    now = datetime.now()
    return render_template("team_leader_objectives_overview.html",batch=batch,employee=employee,objectives=objectives,total_weighted=total_weighted,role="team_leader",state="team_leader",objective=objective,auth=auth,mode=mode,final_total_weighted=final_total_weighted,active=active,now=now)

@team_leader_bp.route("/objective_overview/<int:objective_id>/<int:assigned_by_id>", methods=["POST","GET"])
@login_required
@team_leader_required
def objective_overview(objective_id, assigned_by_id):
    if request.method == "POST":
        message = request.form.get("message")
        if message:
            now = datetime.now()
            new_message = Messages(message=message,status="team_leader",timestamp=now,admin_objective_id=objective_id)
            db.session.add(new_message)
            db.session.commit()
    message = Messages.query.filter_by(admin_objective_id=objective_id).delete()
    db.session.commit()
    messages = Messages.query.filter_by(admin_objective_id=objective_id).order_by(Messages.timestamp.asc()).all()
    assigned_by_auth = Authentication.query.get(assigned_by_id)
    if assigned_by_auth.role == "team_leader":
        assigned_by = assigned_by_auth.team_leader
        objective = Objective.query.get(objective_id)
        batch = objective.batch
        mode = "t"
    elif assigned_by_auth.role == "admin":
        assigned_by = assigned_by_auth.administrator
        objective = AdminObjective.query.get(objective_id)
        batch = objective.batch
        mode = "a"
    else:
        abort(404)
    title = batch.title
    year = batch.year
    employee = objective.assigned_to.name
    confirmed = batch.completed
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()


    auth = Authentication.query.get(session.get("user_id"))
    return render_template(
        "team_leader_objective_overview.html",
        year=year,
        title=title,
        objective=objective,
        employee=employee,
        role="team_leader",
        state="team_leader",
        messages=messages,
        mode=mode,
        confirmed=confirmed,
        batch=batch,
        now=now,
        auth = Authentication.query.get(session.get("user_id"))
    )


@team_leader_bp.route("/delete_objective/<int:objective_id>/<int:assigned_by_id>", methods=["POST", "GET"])
@login_required
@team_leader_required
def delete_objective(objective_id,assigned_by_id):
    assigned_by_auth = Authentication.query.get(assigned_by_id)
    auth_id = assigned_by_auth.id
    if assigned_by_auth.role == "team_leader":
        assigned_by = assigned_by_auth.team_leader
        objective = Objective.query.get(objective_id)
    elif assigned_by_auth.role == "admin":
        assigned_by = assigned_by_auth.administrator
        objective = AdminObjective.query.get(objective_id)
    else:
        abort(404)
    if request.method == "POST":
        db.session.delete(objective)
        db.session.commit()
        flash(f"Objective '{objective.objective}' deleted successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives",auth_id=auth_id,auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    active = objective.batch.active
    now = datetime.now()
    return render_template("delete_objective.html", objective=objective, role="team_leader", state="team_leader",auth=auth,active=active,now=now)

@team_leader_bp.route("/feedback/<int:objective_id>", methods=["GET", "POST"])
@login_required
@team_leader_required
def feedback(objective_id):
    admin_objective = AdminObjective.query.get(objective_id)
    now = datetime.now()
    if request.method == "POST":
        feedback_text = request.form.get("feedback")

        if not feedback_text:
            flash("Feedback is required", "error")
            return redirect(request.url)

        fb = TeamLeaderFeedback(feedback=feedback_text)
        admin_objective.admin_review.team_leader_feedback = fb
        db.session.commit()
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objective_overview", objective_id=admin_objective.id,assigned_by_id=admin_objective.assigned_by_id,auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    active = admin_objective.batch.active
    return render_template(
        "team_leader_feedback.html",
        objective=admin_objective,
        role="team_leader",
        state="team_leader",
        Title="Feedback",
        auth=auth,
        active=active,
        now = datetime.now()
    )

@team_leader_bp.route("/edit_feedback/<int:objective_id>", methods=["GET", "POST"])
@login_required
@team_leader_required
def edit_feedback(objective_id):
    objective = AdminObjective.query.get(objective_id)
    auth_id = objective.assigned_to_id

    
    if objective.assigned_to.id != session["user_id"]:
        flash("Unauthorized", "error")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives",auth_id=auth_id,auth=auth))

    admin_review = objective.admin_review
    if not admin_review:
        flash("No admin review yet", "error")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objectives",auth_id=auth_id,auth=auth))

    if request.method == "POST":
        feedback_text = request.form.get("feedback")
        if not feedback_text:
            flash("Feedback is required", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(request.url,auth=auth)

        if not admin_review.team_leader_feedback:
            admin_review.team_leader_feedback = TeamLeaderFeedback(
                feedback=feedback_text
            )
        else:
            admin_review.team_leader_feedback.feedback = feedback_text

        db.session.commit()
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("team_leader.objective_overview", objective_id=objective.id,assigned_by_id=objective.assigned_by_id,auth=auth))

    feedback_text = (
        admin_review.team_leader_feedback.feedback
        if admin_review.team_leader_feedback
        else ""
    )
    auth = Authentication.query.get(session.get("user_id"))
    active = objective.batch.active
    now = datetime.now()
    return render_template(
        "team_leader_edit_feedback.html",
        objective=objective,
        feedback=feedback_text,
        role="team_leader",
        state="team_leader",
        auth=auth,
        active=active,
        now = datetime.now()
    )