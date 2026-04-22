from flask import Blueprint,url_for,render_template,redirect,flash,get_flashed_messages,request,session,abort
from app.models import Authentication,AdminObjective,Employee,Review,ObjectiveBatch,AdminReview,Administrator,Objective,EmployeeEmail,TeamLeader,Messages,Review,AuthReviewed,ReviewOpenObjective
from app import db
from sqlalchemy.exc import IntegrityError
from app.utils import login_required,admin_required
from collections import defaultdict
from datetime import datetime, timedelta
from docx import Document
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


admin_bp = Blueprint("admin",__name__,url_prefix="/admin")

@admin_bp.route("/",methods=["GET"])
@login_required
@admin_required
def home():
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("home.html",state="admin",role="admin",auth_name=auth_name,auth=auth)

@admin_bp.route("/logout",methods=["POST","GET"])
@login_required
@admin_required
def logout():
    if request.method == "POST":
        session.clear()
        return redirect(url_for("base.login"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("logout.html",state="admin",role="admin",auth=auth)

@admin_bp.route("/delete_account",methods=["POST","GET"])
@login_required
@admin_required
def delete_account():
    if request.method == "POST":
        authen = Authentication.query.get(session["user_id"])
        session.clear()
        if authen:
            db.session.delete(authen.administrator)
            db.session.delete(authen)
        try:
            db.session.commit()
        except IntegrityError as e:
            db.session.rollback()
            flash("Cannot delete account. It is linked to other records.", "error")
            return redirect(url_for("base.admin_signup"))
        flash("account deleted succesfully","success")
        return redirect(url_for("base.login"))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("delete_account.html",state="admin",role="admin",auth=auth)

@admin_bp.route("/add_recipients/<int:batch_id>")
@login_required
@admin_required
def recipients(batch_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    existing_ids = db.session.query(AdminObjective.assigned_to_id).filter_by(batch_id=batch_id).subquery()
    names_auth = Authentication.query.filter(Authentication.id.notin_(existing_ids),Authentication.id!=current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    batch = ObjectiveBatch.query.get(batch_id)
    return render_template("admin_recipients.html",names_auth=names_auth,auth=auth,role="admin",batch_id=batch_id)

@admin_bp.route("/select_member")
@login_required
@admin_required
def select_member():
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_select_member.html",names_auth=names_auth,admin_id=current_user_id,auth=auth,role="admin",state="admin")

@admin_bp.route("/select_batch",methods=["POST","GET"])
@login_required
@admin_required
def select_batch():
    batches = ObjectiveBatch.query.all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("select_batch.html",auth=auth,role="admin",batches=batches,state="admin")

@admin_bp.route("/choose_batch",methods=["POST","GET"])
@login_required
@admin_required
def select_batch_for_open_batch():
    batches = ObjectiveBatch.query.all()
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("select_batch_for_open_batch.html",auth=auth,role="admin",batches=batches,state="admin")


@admin_bp.route("/reports/<auth_id>/<batch_id>",methods=["POST","GET"])
@login_required
@admin_required
def reports(auth_id,batch_id):
    auth = Authentication.query.get(session.get("user_id"))
    grouped = defaultdict(list)
    reports = defaultdict(list)
    batch = ObjectiveBatch.query.get(batch_id)
    objectives = AdminObjective.query.filter_by(batch_id=batch_id).all() + Objective.query.filter_by(batch_id=batch_id).all()
    auths = Authentication.query.filter(Authentication.id != auth_id).all()
    grouped_by_score = defaultdict(int)
    for auth in auths:
        open_total_weighted = 0
        objs = AdminObjective.query.filter_by(batch_id=batch_id, assigned_to_id=auth.id).all()
        if objs:
            for obj in objs:
                if obj.open_objectives_review:
                    open_total_weighted += obj.open_objectives_review.weighted_score
        
        grouped_by_score[auth.id] += open_total_weighted
    for obj in objectives:
        total_weighted = 0
        if hasattr(obj, 'admin_review') and obj.admin_review:
            total_weighted += obj.admin_review.weighted_score
        else:
            if hasattr(obj, 'review') and obj.review:
                total_weighted += obj.review.weighted_score
        if total_weighted:
            key = (obj.batch.title, obj.batch.year, obj.assigned_to)
            op_total_weighted = total_weighted + grouped_by_score[obj.assigned_to_id]
            grouped[key].append(op_total_weighted)
    for key, op_total_weighted in grouped.items():
        title, year, assigned_to = key
        obj = Objective.query.filter_by(batch_id=batch_id,assigned_to_id=assigned_to.id).first()
        if obj:
            obj_id = obj.id
            mode = "t"
        else:
            obj_id = AdminObjective.query.filter_by(batch_id=batch_id,assigned_to_id=assigned_to.id).first().id
            mode = "a"
        name_obj = (assigned_to,obj_id,mode)
        reports[name_obj].append(sum(op_total_weighted))
    return render_template("reports.html",auth=auth,role="admin",state="admin",batch=batch,reports=reports)

@admin_bp.route("/download-report-word/<int:objective_id>")
def download_report_word(objective_id):
    doc = Document()
    doc.add_heading("Objective Report", 0)
    doc.add_paragraph(f"Objective ID: {objective_id}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # add a table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Message"
    hdr[1].text = "Status"
    hdr[2].text = "Timestamp"

    # add data rows
    messages = Messages.query.filter_by(objective_id=objective_id).all()
    for msg in messages:
        row = table.add_row().cells
        row[0].text = msg.message
        row[1].text = msg.status
        row[2].text = str(msg.timestamp)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@admin_bp.route("/download-report/<int:objective_id>")
def download_report(objective_id):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # write content
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, height - 50, "Objective Report")

    p.setFont("Helvetica", 12)
    p.drawString(50, height - 80, f"Objective ID: {objective_id}")
    p.drawString(50, height - 100, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p.save()
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")

@admin_bp.route("/objective_batches",methods=["POST","GET"])
@login_required
@admin_required
def objective_batches():
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    batches = ObjectiveBatch.query.order_by(ObjectiveBatch.year.desc()).all()
    now = datetime.now()
    return render_template("objective_batches.html",auth=auth,role="admin",state="admin",names_auth=names_auth,batches=batches,now=now)

@admin_bp.route("/add_objective_batch",methods=["POST","GET"])
@login_required
@admin_required
def add_objective_batch():
    auth_name = Authentication.query.get(session["user_id"]).name
    current_user_id = session["user_id"]
    names_auth = Authentication.query.filter(Authentication.id != current_user_id).order_by(Authentication.name.asc()).all()
    auth = Authentication.query.get(session.get("user_id"))
    if request.method == "POST":
        year = request.form.get("year")
        title = request.form.get("title")
        duration = request.form.get("duration")
        now = datetime.now()
        deadline = now + timedelta(weeks=int(duration))
        if ObjectiveBatch.query.filter_by(title=title).first():
            flash(f"Batch {title} already exists ","error")
            return redirect(url_for("admin.objective_batches",auth=auth,role="admin",state="admin",names_auth=names_auth))
        new_batch = ObjectiveBatch(title=title, year=year, completed=True, created_at=now, deadline=deadline, duration=duration)
        if new_batch:
            db.session.add(new_batch)
            db.session.commit()
            flash("Batch created successfully","success")
            return redirect(url_for("admin.objective_batches",auth=auth,role="admin",state="admin",names_auth=names_auth))
    return render_template("add_objective_batch.html",auth=auth,role="admin",state="admin",names_auth=names_auth)


@admin_bp.route("/add_objectives/<int:batch_id>",methods=["POST","GET"])
@login_required
@admin_required
def add_objective(batch_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    authen = Authentication.query.get(session["user_id"])
    administrator = authen.administrator if authen else None
    recipients = []
    if request.method == "GET":
        recipient_ids = request.args.getlist("recipients[]")
        recipients.extend(Authentication.query.filter(Authentication.id.in_(recipient_ids)).all())
        if not recipients:
            recipient_ids = request.args.get("recipient_id", type=int)
            recipients.append(Authentication.query.get(recipient_ids))
        auth = Authentication.query.get(session.get("user_id"))
        batch = ObjectiveBatch.query.get(batch_id)
        return render_template("admin_add_objective.html", recipients=recipients,state="admin",auth=auth,role="admin",batch=batch)



    if request.method == "POST":
        recipient_ids = request.form.getlist("recipients[]")
        if recipient_ids:
            recipients.extend(Authentication.query.filter(Authentication.id.in_(recipient_ids)).all())
        if not recipients:
            recipient_ids = request.args.get("recipient_id", type=int)
            recipients.extend(Authentication.query.get(recipient_ids))

        objectives = request.form.getlist("objectives[]")
        categories = request.form.getlist("categories[]")
        weights = request.form.getlist("weights[]")
        publics = request.form.getlist("public[]")
        score_range = 5

       
        batch = ObjectiveBatch.query.get(batch_id)

        
        for obj, cat, weight, public in zip(
            objectives, categories, weights, publics
        ):
            if public == "True":
                private = False
            else:
                private = True
            for recipient in recipients:
                objective = AdminObjective(
                    objective=obj,
                    category=cat,
                    weight=int(weight),
                    score_range=int(score_range),
                    assigned_to=recipient,          
                    assigned_by=administrator,
                    batch=batch,
                    private=private
                )
                db.session.add(objective)
        db.session.commit()
        flash("Objective(s) created successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.objectives",auth_id=recipients[0].id,auth=auth,batch_id=batch.id))


@admin_bp.route("/add_member", methods=["POST", "GET"])
@login_required
@admin_required
def add_member():
    auth_name = Authentication.query.get(session["user_id"]).name
    if request.method == "POST":
        email = request.form.get("email")
        role = request.form.get("role")
        department = request.form.get("department")
        replace = request.form.get("replace")

        existing_email = EmployeeEmail.query.filter_by(email=email).first()

        if existing_email and existing_email.department != department:
            flash("This email is already enrolled in another department.", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

        if existing_email and existing_email.role == role and existing_email.department == department:
            if role == "team_leader":
                role = "team leader"
            flash(f"This email is already enrolled as {role} in this department.", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")
        if existing_email and existing_email.role == "team_leader" and existing_email.department == department:
            role = "team leader"
            flash(f"This email is already enrolled as {role} in this department.", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

        if role == "team_leader":
            existing_tl = EmployeeEmail.query.filter_by(
                role="team_leader", department=department
            ).first()

            if existing_tl and existing_tl.email != email and replace != "yes":
                auth = Authentication.query.get(session.get("user_id"))
                return render_template(
                    "admin_add_member.html",
                    state="admin",
                    confirm_replace=True,
                    existing_tl_email=existing_tl.email,
                    form_email=email,
                    form_role=role,
                    form_department=department,
                    auth=auth,role="admin"
                )

            if existing_tl and existing_tl.email != email and replace == "yes":
                existing_tl.role = "employee"

                old_auth = Authentication.query.filter_by(email=existing_tl.email).first()
                if old_auth and old_auth.team_leader:
                    old_tl = old_auth.team_leader
                    dept = old_tl.department
                    new_emp = Employee(name=old_auth.name, department=dept)
                    old_auth.role = "employee"
                    old_auth.employee = new_emp
                    db.session.add(new_emp)
                    db.session.delete(old_tl)

                db.session.flush()

            if existing_email and existing_email.role == "employee":
                existing_email.role = "team_leader"

                emp_auth = Authentication.query.filter_by(email=email).first()
                if emp_auth and emp_auth.employee:
                    old_emp = emp_auth.employee
                    dept = old_emp.department
                    new_tl = TeamLeader(name=emp_auth.name, department=dept)
                    emp_auth.role = "team_leader"
                    emp_auth.team_leader = new_tl
                    db.session.add(new_tl)
                    db.session.delete(old_emp)

                db.session.commit()
                flash(f"{email} has been promoted to team leader.", "success")
                auth = Authentication.query.get(session.get("user_id"))
                return redirect(url_for("admin.home",auth=auth))

        if not existing_email:
            new_entry = EmployeeEmail(email=email, role=role, department=department)
            db.session.add(new_entry)
            db.session.commit()
            flash("Member enrolled successfully.", "success")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(url_for("admin.home",auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_add_member.html", state="admin",auth=auth,role="admin")

@admin_bp.route("/edit_objectives/<int:objective_id>", methods=["POST", "GET"])
@login_required
@admin_required
def edit_objective(objective_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    admin_email = session.get("email")
    auth = Authentication.query.filter_by(email=admin_email).first()
    administrator = auth.administrator
    
    objective = AdminObjective.query.filter_by(id=objective_id).first()
    if request.method == "POST":
        admin_objective = AdminObjective.query.filter_by(id=objective_id).first()
        objective = request.form.get("objective")
        category = request.form.get("category")
        weight = request.form.get("weight")
        score_range = 5
        emp_name = request.form.get("assigned_to")
        public = request.form.get("public")
        obj = AdminObjective.query.get(objective_id).first()
        batch = obj.batch
        if public:
            private = False
        else:
            private = True
        assigned_to = admin_objective.assigned_to
        
        obj.objective = objective
        obj.category = category
        obj.weight = weight
        obj.private = private
        db.session.commit()
        flash(f"Objective {batch.title} edited successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.objectives",auth_id=assigned_to.id,auth=auth))
    
    authen = Authentication.query.get(session["user_id"])
    administrator = authen.administrator
    departments = administrator.departments
    
    employee_names = []
    team_leader_names = []
    
    for department in departments:
        for emp in department.employees:
            if emp.authentication:
                employee_names.append(emp.authentication.name)
        if department.team_leader and department.team_leader.authentication:
            team_leader_names.append(department.team_leader.authentication.name)
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_edit_objective.html",employee_names=employee_names,objective=objective,team_leader_names=team_leader_names,Title="EDIT OBJECTIVES",state='admin',auth=auth,role="admin")

@admin_bp.route("/open_objectives/<int:batch_id>",methods=["POST","GET"])
@login_required
@admin_required
def open_objectives(batch_id):
    auth = Authentication.query.get(session.get("user_id"))
    objs = defaultdict(list)
    open_objectives = AdminObjective.query.filter(AdminObjective.batch_id == batch_id,AdminObjective.private != True,AdminObjective.assigned_to_id != auth.id).all() + Objective.query.filter(Objective.batch_id == batch_id,Objective.private != True,Objective.assigned_to_id != auth.id).all()
    now = datetime.now()
    for open_objective in open_objectives:
        key = (open_objective.assigned_to,open_objective.batch)
        objs[key] = open_objective

    open_objective = []
    for key,obj in objs.items():
        if isinstance(obj, Objective):
            mode = "t"
        else:
            mode = "a"
        open_objective.append((obj, mode))
    batch = ObjectiveBatch.query.get(batch_id)
    return render_template("open_objectives.html",open_objectives=open_objective,state="admin",role="admin",batch=batch,now=now)


@admin_bp.route("/open_objectives_overview/<int:batch_id>/<int:objective_id>/<mode>",methods=["POST","GET"])
@login_required
@admin_required
def open_objectives_overview(batch_id,objective_id,mode):
    authr = AuthReviewed.query.all()
    auth = Authentication.query.get(session["user_id"])
    auth_name = auth.name
    auth_id = auth.id

    if mode == "a":
        auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=objective_id,auth_id=auth_id).all()
        objective = AdminObjective.query.get(objective_id)
        batch = objective.batch
        employee = objective.assigned_to
        objectives = AdminObjective.query.filter_by(batch_id=batch_id, assigned_to_id=employee.id, private=False).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
        objective = Objective.query.get(objective_id)
        batch = objective.batch
        employee = objective.assigned_to
        objectives = Objective.query.filter_by(batch_id=batch_id, assigned_to_id=employee.id, private=False).all()
    total_weighted = 0
    objs = {}
    for obj in objectives:
        if mode == "a":
            auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=obj.id,auth_id=auth_id).first()
        else:
            auth_reviewed = AuthReviewed.query.filter_by(objective_id=obj.id,auth_id=auth_id).first()
        if hasattr(obj, "open_objectives_review"):  
            if obj.open_objectives_review:
                key = (obj,auth_reviewed)
                objs[key] = obj.open_objectives_review.weighted_score
                total_weighted += obj.open_objectives_review.weighted_score
        
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    return render_template("open_objectives_overview.html",employee=employee,objectives=objs,objective=objectives,total_weighted=total_weighted,batch=batch,state='admin',auth=auth,role="admin",now=now,mode=mode,auth_reviewed=auth_reviewed)

@admin_bp.route("/open_objective_overview/<int:objective_id>/<int:assigned_by_id>/<mode>", methods=["POST","GET"])
@login_required
@admin_required
def open_objective_overview(objective_id, assigned_by_id, mode):
    auth_id = session.get("user_id")
    if mode == "t":
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=objective_id,auth_id=auth_id).all()
    if request.method == "POST":
        message = request.form.get("message")
        if message:
            now = datetime.now()
            new_message = Messages(message=message,status="admin",timestamp=now,admin_objective_id=objective_id)
            db.session.add(new_message)
            db.session.commit()
    db.session.commit()
    messages = Messages.query.filter_by(admin_objective_id=objective_id).order_by(Messages.timestamp.asc()).all()
    assigned_by_auth = Authentication.query.get(assigned_by_id)
    if assigned_by_auth.role == "team_leader":
        assigned_by = assigned_by_auth.team_leader
        objective = Objective.query.get(objective_id)
        batch = objective.batch
    elif assigned_by_auth.role == "admin":
        assigned_by = assigned_by_auth.administrator
        objective = AdminObjective.query.get(objective_id)
        batch = objective.batch
    else:
        abort(404)
    title = batch.title
    year = batch.year
    employee = objective.assigned_to.name
    auth = Authentication.query.get(session.get("user_id"))
    return render_template(
        "open_objective_overview.html",
        year=year,
        title=title,
        objective=objective,
        employee=employee,
        role="admin",
        state="admin",
        messages=messages,
        mode=mode,
        auth = Authentication.query.get(session.get("user_id")),
        auth_reviewed=auth_reviewed)


@admin_bp.route("/review_open_objective/<int:objective_id>/<mode>", methods=["POST", "GET"])
@login_required
@admin_required
def review_open_objective(objective_id,mode):
    auth_id = session.get("user_id")
    if mode == "t":
        auth_reviewed = AuthReviewed.query.filter_by(objective_id=objective_id,auth_id=auth_id).all()
    else:
        auth_reviewed = AuthReviewed.query.filter_by(admin_objective_id=objective_id,auth_id=auth_id).all()
    objective = Objective.query.get(objective_id) if mode == "t" else AdminObjective.query.get(objective_id)
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            auth = Authentication.query.get(session.get("user_id"))
            return redirect(request.url,auth=auth)
        weighted_score = (score * objective.weight) / objective.score_range
        obj = objective.open_objectives_review
        batch = objective.batch
        if obj and obj.weighted_score:
            weighted_scores = obj.weighted_score
        else:
            weighted_scores = 0
        
        if obj and obj.number_reviews:
            number = obj.number_reviews
        else:
            number = 0
        new_number = number + 1
        new_weighted_score = ((weighted_scores * number) + weighted_score)/new_number
        if obj:
            obj.weighted_score = new_weighted_score
            obj.number_reviews = new_number
            obj.score = score
            obj.review = review_text
        else:
            obj = ReviewOpenObjective(
                review=review_text,
                score=score,
                weighted_score=new_weighted_score,
                number_reviews=new_number)
            objective.open_objectives_review = obj
        if isinstance(objective, Objective):
            auth_review = AuthReviewed(
                score=score,
                auth_id=auth_id,
                objective_id = objective_id)
        else:
            auth_review = AuthReviewed(
                score=score,
                auth_id=auth_id,
                admin_objective_id = objective_id)
        db.session.add(auth_review)
        db.session.commit()
        flash("reviewed successfully","success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.open_objectives_overview", objective_id=objective.id, auth=auth, mode=mode, batch_id=batch.id))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_review.html", objective=objective, role="admin", state="admin",auth=auth, mode=mode, auth_reviewed=auth_reviewed)

@admin_bp.route("/objectives/<int:auth_id>/<int:batch_id>")
@login_required
@admin_required
def objectives(auth_id,batch_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    auth = Authentication.query.get(auth_id)
    admin_auth =  Authentication.query.get(session.get("user_id"))
    mode = "See All"
    grouped = defaultdict(list)
    batch = ObjectiveBatch.query.get(batch_id)
    title = batch.title
    objectives = AdminObjective.query.filter_by(batch=batch).all()
    for obj in objectives:
        key = (obj.batch.title, obj.assigned_to_id)
        grouped[key].append(obj)
    username = auth.name
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    return render_template("admin_objectives.html",grouped_objectives=grouped,state='admin',name=username,auth=auth,role="admin",batch_id=batch_id,mode="See All",title=title,batch=batch,now=now)

@admin_bp.route("/review/<int:objective_id>", methods=["POST", "GET"])
@login_required
@admin_required
def review_objective(objective_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    objective = AdminObjective.query.get(objective_id)
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            return redirect(request.url)
        weighted_score = (score * objective.weight) / objective.score_range
        r = AdminReview(
            review=review_text,
            score=score,
            weighted_score=weighted_score)
        objective.admin_review= r
        db.session.commit()
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.objective_overview", objective_id=objective.id,auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_review.html", objective=objective,state='admin',auth=auth,role="admin")

@admin_bp.route("/edit_review/<int:objective_id>", methods=["POST", "GET"])
@login_required
@admin_required
def edit_review(objective_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    objective = AdminObjective.query.get(objective_id)
    review = objective.admin_review
    if request.method == "POST":
        review_text = request.form.get("review")
        score_raw = request.form.get("score")
        if not all([review_text, score_raw]):
            flash("Both Review and Score are needed", "error")
            return redirect(request.url)
        try:
            score = round(float(score_raw), 1)
        except ValueError:
            flash("Score must be a number", "error")
            return redirect(request.url)
        weighted_score = (score * objective.weight) / objective.score_range

        review.review = review_text
        review.score = score
        review.weighted_score = weighted_score
        db.session.commit()
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.objective_overview", objective_id=objective.id,auth=auth))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_edit_review.html", objective=objective,review=review, role="admin", state="admin",Title="Edit Review",auth=auth)


@admin_bp.route("/objectives_overview/<mode>/<int:objective_id>")
@login_required
@admin_required
def objectives_overview(mode,objective_id):
    total_weighted = 0
    final_total_weighted = 0
    if mode == "a":
        objective = AdminObjective.query.get(objective_id)
        employee = objective.assigned_to
        batch = objective.batch
        objectives = AdminObjective.query.filter_by(batch_id=batch.id, assigned_to_id=employee.id).all()
        for obj in objectives:
            if obj.admin_review:
                total_weighted += obj.admin_review.weighted_score
                final_total_weighted += obj.admin_review.weighted_score
                auth = Authentication.query.get(session.get("user_id"))
    else:
        objective = Objective.query.get(objective_id)   
        employee = objective.assigned_to
        batch = objective.batch
        objectives = Objective.query.filter_by(batch_id=batch.id, assigned_to_id=employee.id).all()
        for obj in objectives:
            if obj.review:
                total_weighted += obj.review.weighted_score
                final_total_weighted += obj.admin_review.weighted_score
    for obj in objectives:
        if obj.open_objectives_review:
            final_total_weighted += obj.open_objectives_review.weighted_score
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("admin_objectives_overview.html",batch=batch,employee=employee,objectives=objectives,total_weighted=total_weighted,role="admin",state="admin",objective=objective,auth=auth,mode=mode,final_total_weighted=final_total_weighted)



@admin_bp.route("/objective_overview/<int:objective_id>", methods=["POST", "GET"])
@login_required
@admin_required
def objective_overview(objective_id):
    if request.method == "POST":
        message = request.form.get("message")
        if message:
            now = datetime.now()
            new_message = Messages(message=message,status="admin",timestamp=now,admin_objective_id=objective_id)
            db.session.add(new_message)
            db.session.commit()
    messages = Messages.query.filter_by(admin_objective_id=objective_id).order_by(Messages.timestamp.asc()).all()
    auth_name = Authentication.query.get(session["user_id"]).name
    objective = AdminObjective.query.get(objective_id)
    batch = objective.batch
    title = batch.title
    year = batch.year
    confirmed = batch.completed
    employee = objective.assigned_to.name
    auth = Authentication.query.get(session.get("user_id"))
    now = datetime.now()
    return render_template("admin_objective_overview.html",year=year,title=title,confirmed=confirmed,employee=employee,objective=objective,state='admin',auth=auth,role="admin",messages=messages,batch=batch,now=now)


@admin_bp.route("/delete_objective/<int:objective_id>", methods=["POST", "GET"])
@login_required
@admin_required
def delete_objective(objective_id):
    auth_name = Authentication.query.get(session["user_id"]).name
    objective = AdminObjective.query.get(objective_id)
    batch = objective.batch
    employee = objective.assigned_to
    objectives = AdminObjective.query.filter_by(batch_id=batch.id, assigned_to_id=employee.id).all()
    if request.method == "POST":
        for obj in objectives:
            db.session.delete(obj)
        db.session.commit()
        flash(f"Objective '{objective.objective}' deleted successfully", "success")
        auth = Authentication.query.get(session.get("user_id"))
        return redirect(url_for("admin.objectives",auth_id=auth.id,batch_id=batch.id))
    auth = Authentication.query.get(session.get("user_id"))
    return render_template("delete_objective.html", objective=objective,state='admin',auth=auth,role="admin")